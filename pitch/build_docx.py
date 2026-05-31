# -*- coding: utf-8 -*-
"""Build a polished RTL Hebrew Word doc of the meeting-prep material."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BOLT = RGBColor(0x5c, 0x33, 0x17)
DARK = RGBColor(0x4a, 0x2a, 0x13)
HOT = RGBColor(0xc9, 0x62, 0x2f)
UP = RGBColor(0xb4, 0x47, 0x2d)
DOWN = RGBColor(0x4a, 0x7c, 0x3f)
SUB = RGBColor(0x8a, 0x6a, 0x4a)
FONT = "Arial"


def _rtl_para(p):
    pPr = p._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    pPr.append(bidi)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def _cs(run):
    """Mark run text as complex-script so Hebrew uses the cs font/size."""
    rPr = run._element.get_or_add_rPr()
    rfonts = rPr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rPr.append(rfonts)
    rfonts.set(qn("w:cs"), FONT)
    rfonts.set(qn("w:ascii"), FONT)
    rfonts.set(qn("w:hAnsi"), FONT)


def heading(doc, text, level=1, color=BOLT, size=None):
    p = doc.add_paragraph()
    _rtl_para(p)
    r = p.add_run(text)
    r.bold = True
    r.font.name = FONT
    r.font.color.rgb = color
    r.font.size = Pt(size or (20 if level == 1 else 14))
    _cs(r)
    if level == 1:
        # bottom border
        pPr = p._p.get_or_add_pPr()
        pbdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "8")
        bottom.set(qn("w:space"), "4")
        bottom.set(qn("w:color"), "C9622F")
        pbdr.append(bottom)
        pPr.append(pbdr)
    p.space_after = Pt(6)
    return p


def body(doc, text, bold=False, color=None, size=11, italic=False, space=4):
    p = doc.add_paragraph()
    _rtl_para(p)
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.name = FONT
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color
    _cs(r)
    p.paragraph_format.space_after = Pt(space)
    return p


def bullet(doc, text, color=None):
    p = doc.add_paragraph(style="List Bullet")
    _rtl_para(p)
    r = p.add_run(text)
    r.font.name = FONT
    r.font.size = Pt(11)
    if color:
        r.font.color.rgb = color
    _cs(r)
    return p


def quote(doc, text):
    p = doc.add_paragraph()
    _rtl_para(p)
    pPr = p._p.get_or_add_pPr()
    # right border bar
    pbdr = OxmlElement("w:pBdr")
    right = OxmlElement("w:right")
    right.set(qn("w:val"), "single")
    right.set(qn("w:sz"), "18")
    right.set(qn("w:space"), "8")
    right.set(qn("w:color"), "5C3317")
    pbdr.append(right)
    pPr.append(pbdr)
    p.paragraph_format.right_indent = Pt(6)
    p.paragraph_format.left_indent = Pt(12)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    r.italic = True
    r.font.name = FONT
    r.font.size = Pt(11.5)
    r.font.color.rgb = DARK
    _cs(r)
    return p


def make_table(doc, rows, header=True):
    t = doc.add_table(rows=0, cols=len(rows[0]))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.RIGHT
    # RTL visual order
    tblPr = t._tbl.tblPr
    bidi = OxmlElement("w:bidiVisual")
    tblPr.append(bidi)
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for ci, val in enumerate(row):
            cell = cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            _rtl_para(p)
            r = p.add_run(val)
            r.font.name = FONT
            r.font.size = Pt(10.5)
            _cs(r)
            if ri == 0 and header:
                r.bold = True
                r.font.color.rgb = RGBColor(0xff, 0xff, 0xff)
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:fill"), "5C3317")
                cell._tc.get_or_add_tcPr().append(shd)
    return t


def spacer(doc):
    doc.add_paragraph()


doc = Document()
style = doc.styles["Normal"]
style.font.name = FONT
style.font.size = Pt(11)

# ── Title ─────────────────────────────────────────────────────────
p = doc.add_paragraph()
_rtl_para(p)
r = p.add_run("MOCA — ערכת הכנה לפגישה")
r.bold = True
r.font.size = Pt(26)
r.font.name = FONT
r.font.color.rgb = BOLT
_cs(r)
body(doc, "פגישה עם סמנכ\"ל שיווק, חברת סלולר גדולה", color=SUB, size=13)
quote(doc, "מטרת הפגישה: לא למכור כלי. למכור שליטה, מהירות תגובה ושקט נפשי. היעד לסיום: לקבל כן לפילוט של שבועיים על 3 מתחרים.")

# ── 0. Numbers ────────────────────────────────────────────────────
heading(doc, "0. מספרים מאומתים מהקוד (לשנן)")
make_table(doc, [
    ["נתון", "מספר", "מקור"],
    ["מפעילים סלולריים מנוטרים", "10", "app.py: CARRIER_DISPLAY"],
    ["חזיתות e-store בצילום", "4", "פלאפון, סלקום, פרטנר, הוט מובייל"],
    ["ספקי eSIM גלובליים", "20+", "24 סקרייפרים ב-scraper.py"],
    ["תדירות סריקה", "פעמיים ביום", "07:30 + 17:00"],
    ["צילום באנרים", "כל בוקר 08:00", "homepage + e-store"],
])
body(doc, "המפעילים: פרטנר · פלאפון · הוט מובייל · סלקום · 019 · XPhone · וי-קום · נפטוקום · גולן · רמי לוי.", space=2)
body(doc, "ספקי eSIM בולטים: Airalo · Holafly · Saily · Orbit · Maya · GoMo · Sparks · VOYE.")
body(doc, "⚠ אם הסמנכ\"ל שואל מספר ואתה לא בטוח — אל תמציא. אמור 'אבדוק ואחזור אליך'. זה בונה אמון.", color=UP, bold=True)

# ── 1. Opening ────────────────────────────────────────────────────
heading(doc, "1. משפט הפתיחה (30 שניות)")
quote(doc, "כמה זמן לוקח לך היום לדעת שמתחרה הוריד מחיר על חבילה — שעות? יום? ומי מביא לך את זה, מישהו שנכנס ידנית לאתרים בבוקר? אני בניתי מערכת שעושה את זה אוטומטית על כל השוק, פעמיים ביום, ושולחת לך התראה תוך דקות.")
body(doc, "עצור. תן לו להגיב. השוק שלו הוא הכאב — תן לו לאשר אותו לפני שאתה ממשיך.", italic=True, color=SUB)

# ── 2. Meeting order ──────────────────────────────────────────────
heading(doc, "2. סדר הפגישה (20–25 דק')")
for t in [
    "הכאב (2 דק') — משפט הפתיחה + הקשבה",
    "המשפט המסכם (1 דק') — מה זה בעצם",
    "הדגמה חיה (8–10 דק') — לב הפגישה. שלושת המסכים",
    "הזווית האסטרטגית (3 דק') — eSIM ונדידה",
    "בידול (2 דק') — למה לא אקסל/חברת מחקר",
    "הסגירה (2 דק') — הצעת הפילוט",
    "שאלות — מענה התנגדויות",
]:
    bullet(doc, t)

# ── 3. Demo script ────────────────────────────────────────────────
heading(doc, "3. סקריפט ההדגמה — צעד אחר צעד")
body(doc, "כלל הזהב: הדגמה מנצחת מצגת. אם אין אינטרנט — צילומי מסך מוכנים מראש.", bold=True, color=HOT)
heading(doc, "הדגמה 1 — מכונת הזמן (/archive)", level=2, color=DARK)
for t in [
    "פתח את המסך. בחר את החברה של הסמנכ\"ל עצמו ('בוא נתחיל ממך').",
    "גלגל לתאריך מלפני ~3 שבועות.",
    "'הנה — כך נראה עמוד הבית שלך, אלה היו החבילות והמחירים.'",
    "החלף למתחרה ראשי. 'ואותו דבר על המתחרה. תיעוד מלא, אוטומטי.'",
]:
    bullet(doc, t)
heading(doc, "הדגמה 2 — קיר הבאנרים (/banners)", level=2, color=DARK)
for t in [
    "הצג את המוזאיקה של צילומי המתחרים.",
    "הצבע על תווית הרעננות: 'היום / אתמול / לפני N ימים'.",
    "'אתה רואה ברגע מי החליף באנר ראשי הבוקר ומה הוא דוחף.'",
]:
    bullet(doc, t)
heading(doc, "הדגמה 3 — Market Movers + התראה", level=2, color=DARK)
for t in [
    "הצג את ווידג'ט 'מי זז הכי הרבה'.",
    "הצבע על הקוד הצבעוני: אדום = מחיר עלה (רע לנו), ירוק = ירד (טוב לנו).",
    "'כל שינוי נשלח לוואטסאפ תוך דקות, עם דדופ חכם שלא מתריע פעמיים.'",
]:
    bullet(doc, t)

# ── 4. Strategic angle ────────────────────────────────────────────
heading(doc, "4. הזווית האסטרטגית — eSIM ונדידה")
quote(doc, "הכנסות הנדידה הן פרה חולבת שנשחקת בדיוק על ידי ספקי ה-eSIM הגלובליים. אני עוקב אחרי 20+ כאלה על מאות מדינות. אתה היחיד בשוק שידע בכמה Airalo מוכר ליוון השבוע מול חבילת הנדידה שלך.")
body(doc, "זו לא עוד טבלה — זו הגנה על שורת הכנסה ספציפית שמדממת אצל כל מפעיל גדול.", bold=True)

# ── 5. Differentiation ────────────────────────────────────────────
heading(doc, "5. הבידול")
quote(doc, "חברת מחקר תיתן לך תמונה של אתמול. אני נותן לך את היום — וגם את ההיסטוריה כדי לראות מגמות.")
make_table(doc, [
    ["קטגוריה", "מה שיש היום", "MOCA"],
    ["תדירות", "ידני, מתי שמתפנים", "אוטומטי, 2×/יום"],
    ["היקף", "מחירים בלבד", "מחיר + קרייטיב + חדשות + שוק אפור + eSIM"],
    ["מהירות תגובה", "שעות–יום", "דקות"],
    ["היסטוריה", "אין", "ארכיון + מכונת זמן"],
    ["תובנות", "נתון גולמי", "ניתוח AI + מיצוב + מגמות"],
])

# ── 6. Objections ─────────────────────────────────────────────────
heading(doc, "6. מענה להתנגדויות")
objs = [
    ("זה חוקי? סקרייפינג של אתרים?",
     "אני אוסף מידע פומבי גלוי — בדיוק מה שכל לקוח רואה כשהוא נכנס לאתר. מודיעין שוק ממקורות גלויים, לא חדירה. מוכן להוריד מתחרה ספציפי אם המשפטיים מבקשים."),
    ("מה הדיוק? מה אם מוצג מחיר שגוי?",
     "לפני הפגישה ריצתי ולידציה על המתחרים שהדגמתי. יש שכבת זיהוי שינויים שמנטרלת התראות-שווא (תנודות מטבע לא נספרות). כל אירוע ניתן לקליק עד צילום המקור."),
    ("למה שלא נבנה לבד / יש לנו צוות BI?",
     "אתם יכולים — בעוד שנה ושני אנשי פיתוח. זה כבר רץ בפרודקשן: תשתית, גיבויים, תזמון, ניטור בריאות. אתם קונים זמן-לשוק, לא קוד. ואני מתחזק את הסקרייפרים כשמתחרה משנה אתר."),
    ("כמה זה עולה?",
     "בוא נתחיל מפילוט בלי מחויבות. אחרי שתראה ערך, נדבר על מודל מנוי חודשי. אל תיתן מספר בפגישה ראשונה אלא אם לוחץ — ואז טווח, לא נקודה."),
    ("איך זה משתלב עם הכלים שלנו?",
     "התראות לאן שתרצו — וואטסאפ, טלגרם, מייל. יש API. ובהמשך אפשר להזרים ל-Slack / Teams / BI."),
]
for q, a in objs:
    body(doc, "❝ " + q, bold=True, color=BOLT, space=1)
    quote(doc, a)

# ── 7. Checklist ──────────────────────────────────────────────────
heading(doc, "7. צ'קליסט לפני הפגישה")
for t in [
    "רוץ scrape ידני על 2-3 המתחרים → ודא עין שהמחירים נכונים.",
    "ודא שהבאנרים של היום נמשכו וטריים.",
    "הכן צילומי מסך מוכנים של שלוש ההדגמות — גיבוי אם אין אינטרנט.",
    "בדוק שמכונת הזמן עובדת על תאריך עבר אמיתי.",
    "גלה מי 3 המתחרים שהכי מטרידים — וכוון אליהם את ההדגמה.",
    "פתח את דף הנחיתה כ-leave-behind, או הדפס.",
    "טען התראה לוואטסאפ מראש כדי להראות התראה אמיתית.",
]:
    bullet(doc, t)

# ── 8. Close ──────────────────────────────────────────────────────
heading(doc, "8. הסגירה")
quote(doc, "תן לי שבועיים. אני מחבר מעקב על שלושת המתחרים שהכי מטרידים אותך ושולח לצוות שלך התראות בזמן אמת. אם בסוף השבועיים זה לא חוסך לכם לפחות ישיבת בוקר אחת — לא שילמת. מתי נוח לך שאתחיל?")
body(doc, "שאלת סגירה אקטיבית ('מתי נוח לך שאתחיל') עדיפה על 'מה דעתך?' — היא מניחה שהתשובה כן.", italic=True, color=SUB)

out = r"d:\\השוואת MASS MARKET\\pitch\\MOCA-meeting-prep.docx"
doc.save(out)
print("SAVED:", out)
